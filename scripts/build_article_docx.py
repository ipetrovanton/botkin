"""Сборка docx-версии статьи для Хабра из Markdown-исходника.

Поддерживается подмножество Markdown, используемое в статье: заголовки,
абзацы с **жирным**/*курсивом*/`кодом`, списки, цитаты, код-блоки,
таблицы в pipe-нотации и мем-блоки (цитаты, начинающиеся с «[МЕМ»).

Запуск: python scripts/build_article_docx.py <in.md> <out.docx>
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

MONO = "Consolas"
CODE_BG = "F2F2F2"
MEME_BG = "FFF3CD"

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*[^*]+?\*|`[^`]+?`)")


def shade(paragraph, color: str) -> None:
    pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), color)
    pr.append(shd)


def add_inline(paragraph, text: str) -> None:
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = MONO
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        elif part.startswith("*") and part.endswith("*"):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        else:
            paragraph.add_run(part)


def add_code_block(doc, lines: list[str]) -> None:
    p = doc.add_paragraph()
    shade(p, CODE_BG)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    for i, line in enumerate(lines):
        run = p.add_run(line)
        run.font.name = MONO
        run.font.size = Pt(8.5)
        if i < len(lines) - 1:
            run.add_break()


def add_table(doc, rows: list[list[str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    for r, row in enumerate(rows):
        for c, cell_text in enumerate(row):
            cell = table.cell(r, c)
            cell.paragraphs[0].text = ""
            add_inline(cell.paragraphs[0], cell_text)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9.5)
                if r == 0:
                    run.bold = True
    doc.add_paragraph()


def is_table_line(line: str) -> bool:
    return line.strip().startswith("|") and line.strip().endswith("|")


def split_row(line: str) -> list[str]:
    return [c.strip() for c in line.strip().strip("|").split("|")]


def build(md_path: Path, out_path: Path) -> None:
    lines = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if stripped.startswith("```"):
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            add_code_block(doc, block)
        elif is_table_line(stripped):
            rows = []
            while i < len(lines) and is_table_line(lines[i].strip()):
                cells = split_row(lines[i])
                if not all(re.fullmatch(r":?-{2,}:?", c) for c in cells):
                    rows.append(cells)
                i += 1
            i -= 1
            if rows:
                add_table(doc, rows)
        elif stripped.startswith("# "):
            p = doc.add_heading(level=0)
            add_inline(p, stripped[2:])
        elif stripped.startswith("## "):
            p = doc.add_heading(level=1)
            add_inline(p, stripped[3:])
        elif stripped.startswith("### "):
            p = doc.add_heading(level=2)
            add_inline(p, stripped[4:])
        elif stripped.startswith("> "):
            quote = stripped[2:]
            p = doc.add_paragraph()
            if quote.startswith("[МЕМ"):
                shade(p, MEME_BG)
                run = p.add_run("🖼 ")
                run.font.size = Pt(11)
                add_inline(p, quote)
                for run in p.runs:
                    run.italic = True
            else:
                p.paragraph_format.left_indent = Pt(24)
                add_inline(p, quote)
                for run in p.runs:
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x40, 0x40, 0x40)
        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, stripped[2:])
        elif re.match(r"^\d+\.\s", stripped):
            p = doc.add_paragraph(style="List Number")
            add_inline(p, re.sub(r"^\d+\.\s", "", stripped))
        elif stripped == "---":
            p = doc.add_paragraph("⸻")
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped:
            p = doc.add_paragraph()
            add_inline(p, stripped)
        i += 1

    doc.save(out_path)
    print(f"OK: {out_path} ({out_path.stat().st_size} bytes)")


if __name__ == "__main__":
    build(Path(sys.argv[1]), Path(sys.argv[2]))
